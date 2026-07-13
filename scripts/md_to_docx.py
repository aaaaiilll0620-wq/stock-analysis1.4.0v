# -*- coding: utf-8 -*-
"""
Markdown → Word (.docx) 轉檔器 — 供開發日誌 / 使用指南等文件產 Word 版
================================================================================
支援本專案文件用到的 GFM 子集:
  標題 #~####、表格 (| ... |)、引用區塊 (>)、圍欄程式碼 (```)、
  有序/無序清單、水平線 (---)、行內 **粗體**、`程式碼`、[連結](url)。
用法:
  python scripts/md_to_docx.py docs/開發日誌_DevLog_135557.md docs/開發日誌_DevLog.docx
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT = "Microsoft JhengHei"
CODE_FONT = "Consolas"


def _set_run(run, bold=False, code=False, size=None, color=None):
    run.font.name = CODE_FONT if code else FONT
    # 中文字型需同時設定 eastAsia,否則 Word 會退回預設中文字體
    rpr = run._element.get_or_add_rPr()
    from docx.oxml.ns import qn
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), CODE_FONT if code else FONT)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _add_inline(par, text, base_size=None, base_color=None):
    """把行內 markdown (粗體/程式碼/連結) 拆成 runs 寫進段落。"""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            r = par.add_run(tok[2:-2])
            _set_run(r, bold=True, size=base_size, color=base_color)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            _set_run(r, code=True, size=base_size or 9.5, color=(0x8B, 0x00, 0x00))
        elif tok.startswith("[") and "](" in tok:
            label = tok[1:tok.index("](")]
            r = par.add_run(label)
            _set_run(r, size=base_size, color=(0x1F, 0x4E, 0x79))
        else:
            r = par.add_run(tok)
            _set_run(r, size=base_size, color=base_color)


def _is_table_sep(line):
    s = line.strip()
    return bool(s.startswith("|")) and set(s.replace("|", "").replace(":", "").strip()) <= {"-", " "} and "-" in s


def _split_row(line):
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def convert(md_path: str, out_path: str) -> None:
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        # 圍欄程式碼區塊
        if s.startswith("```"):
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(lines[i])
                _set_run(r, code=True, size=9, color=(0x40, 0x40, 0x40))
                i += 1
            i += 1
            continue

        # 表格
        if s.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _split_row(s)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.style = "Table Grid"
            for j, cell_text in enumerate(header):
                cell = tbl.rows[0].cells[j]
                cell.text = ""
                _add_inline(cell.paragraphs[0], cell_text, base_size=9.5)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for ri, row in enumerate(rows, start=1):
                for j in range(len(header)):
                    cell = tbl.rows[ri].cells[j]
                    cell.text = ""
                    _add_inline(cell.paragraphs[0], row[j] if j < len(row) else "", base_size=9.5)
            doc.add_paragraph()
            continue

        # 標題
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            level = min(len(m.group(1)), 4)
            p = doc.add_heading("", level=level)
            _add_inline(p, m.group(2))
            for r in p.runs:
                r.font.name = FONT
                from docx.oxml.ns import qn as _qn
                r._element.get_or_add_rPr().find(_qn("w:rFonts")).set(_qn("w:eastAsia"), FONT)
            i += 1
            continue

        # 水平線
        if s in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            from docx.oxml.ns import qn as _qn
            ppr = p._element.get_or_add_pPr()
            pbdr = ppr.makeelement(_qn("w:pBdr"), {})
            bottom = ppr.makeelement(_qn("w:bottom"), {
                _qn("w:val"): "single", _qn("w:sz"): "6",
                _qn("w:space"): "1", _qn("w:color"): "999999"})
            pbdr.append(bottom)
            ppr.append(pbdr)
            i += 1
            continue

        # 引用區塊 (連續 > 行合併,行間保留換行)
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            first = True
            while i < n and lines[i].strip().startswith(">"):
                body = lines[i].strip().lstrip(">").strip()
                if not first:
                    p.add_run().add_break()
                if body:
                    _add_inline(p, body, base_size=9.5, base_color=(0x40, 0x40, 0x40))
                first = False
                i += 1
            continue

        # 清單 (有序/無序;縮排接續行併入同段)
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if m:
            ordered = m.group(2)[0].isdigit()
            p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
            _add_inline(p, m.group(3))
            i += 1
            while i < n and lines[i].startswith(("   ", "\t")) and lines[i].strip() \
                    and not re.match(r"^\s*([-*+]|\d+\.)\s+", lines[i]) \
                    and not lines[i].strip().startswith(("|", "#", ">", "```")):
                p.add_run(" ")
                _add_inline(p, lines[i].strip())
                i += 1
            continue

        # 空行
        if not s:
            i += 1
            continue

        # 一般段落 (連續非空、非結構行合併)
        p = doc.add_paragraph()
        _add_inline(p, s)
        i += 1
        while i < n and lines[i].strip() and not re.match(r"^(\s*([-*+]|\d+\.)\s+|#|\||>|```|---)", lines[i]):
            p.add_run(" ")
            _add_inline(p, lines[i].strip())
            i += 1

    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/開發日誌_DevLog_135557.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "docs/開發日誌_DevLog.docx"
    convert(src, dst)
